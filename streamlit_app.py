"""Streamlit interface for Mariner Advocate."""

from __future__ import annotations

import hashlib
import hmac
import json
import os
import re
import zipfile
from dataclasses import fields
from io import BytesIO
from pathlib import Path, PurePosixPath
from typing import Any
from urllib.parse import urlparse

import streamlit as st
from docx import Document
from pypdf import PdfReader

from mariner_core import (
    AppError,
    MAX_CONTRACT_CHARS,
    ProviderConfig,
    analyze_case,
    answer_followup,
    draft_document,
    extract_contract_fields,
)


st.set_page_config(
    page_title="Mariner Advocate",
    page_icon="⚓",
    layout="wide",
    initial_sidebar_state="expanded",
)


STATIC_CSS = """
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Barlow:wght@400;500;600&family=Barlow+Condensed:wght@500;600;700&display=swap" rel="stylesheet">
<style>
  /* ---- Industry: steel-blue blueprint on a light technical ground ---- */
  :root {
    --bg: #f2f2f3;
    --surface: #e9e9ea;
    --ink: #1d1f20;
    --muted: #5d5d60;
    --faint: #7a7a7d;
    --line: rgba(29, 31, 32, .16);
    --rule: rgba(29, 31, 32, .28);
    --accent: #5980a6;
    --accent-deep: #416180;
    --accent-field: #1d2d3d;
    --heading: "Barlow Condensed", system-ui, sans-serif;
    --body: "Barlow", system-ui, sans-serif;
    /* corner registration marks, drawn as background layers */
    --mk:
      linear-gradient(#1d1f208c,#1d1f208c) left -4px top 0/9px 1px no-repeat,
      linear-gradient(#1d1f208c,#1d1f208c) left 0 top -4px/1px 9px no-repeat,
      linear-gradient(#1d1f208c,#1d1f208c) right -4px top 0/9px 1px no-repeat,
      linear-gradient(#1d1f208c,#1d1f208c) right 0 top -4px/1px 9px no-repeat,
      linear-gradient(#1d1f208c,#1d1f208c) left -4px bottom 0/9px 1px no-repeat,
      linear-gradient(#1d1f208c,#1d1f208c) left 0 bottom -4px/1px 9px no-repeat,
      linear-gradient(#1d1f208c,#1d1f208c) right -4px bottom 0/9px 1px no-repeat,
      linear-gradient(#1d1f208c,#1d1f208c) right 0 bottom -4px/1px 9px no-repeat;
    --mk-light:
      linear-gradient(#f2f2f3a8,#f2f2f3a8) left -4px top 0/9px 1px no-repeat,
      linear-gradient(#f2f2f3a8,#f2f2f3a8) left 0 top -4px/1px 9px no-repeat,
      linear-gradient(#f2f2f3a8,#f2f2f3a8) right -4px top 0/9px 1px no-repeat,
      linear-gradient(#f2f2f3a8,#f2f2f3a8) right 0 top -4px/1px 9px no-repeat,
      linear-gradient(#f2f2f3a8,#f2f2f3a8) left -4px bottom 0/9px 1px no-repeat,
      linear-gradient(#f2f2f3a8,#f2f2f3a8) left 0 bottom -4px/1px 9px no-repeat,
      linear-gradient(#f2f2f3a8,#f2f2f3a8) right -4px bottom 0/9px 1px no-repeat,
      linear-gradient(#f2f2f3a8,#f2f2f3a8) right 0 bottom -4px/1px 9px no-repeat;
  }

  /* ---- ground and type ---- */
  .stApp { background: var(--bg); }
  html, body, .stApp { font-family: var(--body); color: var(--ink); }
  .block-container { max-width: 1500px; padding-top: 1.8rem; padding-bottom: 4rem; }
  h1, h2, h3, h4, h5, h6,
  .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4 {
    font-family: var(--heading); font-weight: 600; letter-spacing: -.01em; line-height: 1.08;
  }
  .stMarkdown h2, .stMarkdown h3, .stMarkdown h4 { text-transform: uppercase; letter-spacing: .02em; }
  .stMarkdown p, .stMarkdown li { font-size: 14.5px; line-height: 1.6; }
  a, a:visited { color: var(--accent-deep); text-underline-offset: 3px; }
  a:hover { color: #2c455d; }
  ::selection { background: #d6ebff; }
  :focus-visible { outline: 2px solid var(--accent); outline-offset: 2px; }
  footer, #MainMenu { visibility: hidden; }

  /* ---- hero: a spec-sheet plate on the paper ground, not a gradient card ---- */
  .hero {
    position: relative;
    margin-bottom: .9rem;
    padding: 1.9rem 2rem;
    border: 1px solid var(--line);
    background: transparent;
    color: var(--ink);
  }
  .hero::before { content: ""; position: absolute; inset: 0; background: var(--mk); pointer-events: none; }
  .hero .eyebrow {
    font-family: var(--heading); font-weight: 600; font-size: .72rem;
    letter-spacing: .16em; text-transform: uppercase; color: var(--accent-deep); opacity: 1;
  }
  .hero h1 {
    font-size: clamp(2.2rem, 3.6vw, 3.5rem); line-height: 1.02;
    margin: .5rem 0 .6rem; max-width: 22ch; letter-spacing: -.02em;
  }
  .hero p { max-width: 74ch; font-size: 1rem; line-height: 1.6; color: #424244; opacity: 1; margin: 0; }

  /* ---- workflow cards: transparent line drawings with corner marks ---- */
  .workflow-card {
    position: relative;
    border: 1px solid var(--line);
    border-radius: 0;
    background: transparent;
    padding: .95rem 1rem;
    min-height: 112px;
  }
  .workflow-card::before { content: ""; position: absolute; inset: 0; background: var(--mk); pointer-events: none; }
  .workflow-card .number {
    font-family: var(--heading); font-weight: 600; font-size: .72rem;
    letter-spacing: .14em; color: var(--accent-deep);
  }
  .workflow-card strong {
    display: block; margin: .3rem 0 .25rem;
    font-family: var(--heading); font-weight: 600; font-size: 1.06rem;
    letter-spacing: .02em; text-transform: uppercase; color: var(--ink);
  }
  .workflow-card small { color: var(--muted); font-size: .8rem; line-height: 1.45; }

  /* ---- the waiting plate ---- */
  .case-empty {
    border: 1px dashed var(--rule);
    border-radius: 0;
    background: transparent;
    padding: 3.4rem 2rem;
    text-align: center;
    color: var(--muted);
  }
  .case-empty h3 {
    color: var(--ink); margin: .6rem 0 .5rem;
    font-family: var(--heading); font-size: 1.9rem; text-transform: uppercase;
  }
  .small-note { color: var(--faint); font-size: .82rem; }

  /* ---- forms: square-cornered, hairline, surface-filled inputs ---- */
  div[data-testid="stForm"] {
    border: 1px solid var(--line); border-radius: 0;
    background: transparent; padding: 1.2rem 1.3rem;
  }
  .stTextInput label, .stTextArea label, .stSelectbox label, .stCheckbox label {
    font-family: var(--body) !important; font-size: 12px !important; color: #424244 !important;
  }
  .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] > div,
  .stNumberInput input {
    background: var(--surface) !important;
    border: 1px solid var(--line) !important;
    border-radius: 4px !important;
    color: var(--ink) !important;
    font-family: var(--body) !important;
  }
  .stTextInput input:focus, .stTextArea textarea:focus { border-color: var(--accent) !important; box-shadow: none !important; }
  .stTextInput input::placeholder, .stTextArea textarea::placeholder { color: #98989b !important; }

  /* ---- buttons: the primary is the one solid object on the board ---- */
  .stButton button, .stDownloadButton button, .stFormSubmitButton button, .stLinkButton a {
    border-radius: 0 !important;
    font-family: var(--heading) !important; font-weight: 600 !important;
    letter-spacing: .07em; text-transform: uppercase;
    border: 1px solid var(--rule) !important;
    background: transparent !important; color: var(--ink) !important;
    transition: background .12s ease;
  }
  .stButton button:hover, .stDownloadButton button:hover, .stLinkButton a:hover {
    background: rgba(29,31,32,.07) !important;
  }
  .stFormSubmitButton button[kind="primaryFormSubmit"],
  .stButton button[kind="primary"] {
    position: relative;
    background: var(--accent) !important; border-color: var(--accent) !important;
    color: var(--bg) !important; padding: .8rem 1.1rem !important; font-size: 1.05rem !important;
  }
  .stFormSubmitButton button[kind="primaryFormSubmit"]::before,
  .stButton button[kind="primary"]::before {
    content: ""; position: absolute; inset: 0; background: var(--mk-light); pointer-events: none;
  }
  .stFormSubmitButton button[kind="primaryFormSubmit"]:hover,
  .stButton button[kind="primary"]:hover { background: #597ea3 !important; border-color: #597ea3 !important; }

  /* ---- containers, expanders, tabs ---- */
  div[data-testid="stVerticalBlockBorderWrapper"] > div:has(> div[data-testid="stVerticalBlock"]) { border-radius: 0; }
  div[data-testid="stExpander"] details {
    border: 1px solid var(--line) !important; border-radius: 0 !important; background: transparent !important;
  }
  div[data-testid="stExpander"] summary p {
    font-family: var(--heading) !important; font-weight: 600; letter-spacing: .05em; text-transform: uppercase;
  }
  .stTabs [data-baseweb="tab-list"] { gap: 0; border-bottom: 1px solid var(--line); }
  .stTabs [data-baseweb="tab"] {
    font-family: var(--heading); font-weight: 600; font-size: .95rem;
    letter-spacing: .08em; text-transform: uppercase; color: var(--faint);
    padding: .7rem .9rem; border-radius: 0;
  }
  .stTabs [aria-selected="true"] { color: var(--ink); }
  .stTabs [data-baseweb="tab-highlight"] { background: var(--accent); }

  /* ---- metrics read as spec-sheet numbers ---- */
  [data-testid="stMetric"] { border: 1px solid var(--line); padding: .8rem .9rem; }
  [data-testid="stMetricLabel"] p {
    font-size: 11px !important; letter-spacing: .12em; text-transform: uppercase; color: var(--faint) !important;
  }
  [data-testid="stMetricValue"] {
    color: var(--accent-deep); font-family: var(--heading); font-weight: 600; font-size: 1.9rem;
  }

  /* ---- callouts: mono palette, one steel edge, no decorative colour ---- */
  div[data-testid="stAlert"] {
    border-radius: 0; border: 1px solid var(--line); border-left: 3px solid var(--accent);
    background: var(--surface); color: #424244;
  }
  div[data-testid="stAlert"] p { font-size: 13.5px; line-height: 1.55; }
  div[data-testid="stAlert"] svg { color: var(--accent-deep); }
  /* the executive summary st.info reads as the one reversed field */
  div[data-testid="stAlert"]:has(svg[title="info"]),
  .exec-field {
    background: var(--accent-field); color: #f2f2f3; border-color: var(--accent-field);
    border-left-color: var(--accent);
  }
  div[data-testid="stAlert"]:has(svg[title="info"]) p { color: #f2f2f3; font-size: 15px; }

  /* ---- sidebar ---- */
  section[data-testid="stSidebar"] {
    background: var(--surface); border-right: 1px solid var(--line);
  }
  section[data-testid="stSidebar"] h2, section[data-testid="stSidebar"] h1 {
    font-family: var(--heading); text-transform: uppercase; letter-spacing: .04em;
  }
  section[data-testid="stSidebar"] .stCode, section[data-testid="stSidebar"] pre {
    background: transparent !important; border: 1px solid var(--line); border-radius: 0;
  }
  hr, div[data-testid="stDivider"] hr { border-color: var(--line); }
  div[data-testid="stStatusWidget"], div[data-testid="stStatus"] {
    border-radius: 0 !important; border: 1px solid var(--line) !important;
  }
</style>
"""


def get_setting(name: str, default: Any = "") -> Any:
    """Read a setting from environment or Streamlit secrets without exposing it."""
    env_value = os.getenv(name)
    if env_value not in (None, ""):
        return env_value
    try:
        if name in st.secrets:
            return st.secrets[name]
    except (FileNotFoundError, KeyError, RuntimeError):
        pass
    return default


def as_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    return str(value or "").strip().lower() in {"1", "true", "yes", "on"}


PUBLIC_DEMO_ONLY = as_bool(get_setting("PUBLIC_DEMO_ONLY", False))
DEMO_MODE = PUBLIC_DEMO_ONLY or as_bool(get_setting("DEMO_MODE", False))
MAX_CONTRACT_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_CONTRACT_PDF_PAGES = 150
MAX_DOCX_ENTRIES = 3_000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
CONTRACT_CASE_FIELDS = {
    "shipName": "Ship name",
    "imoNumber": "IMO number",
    "flagState": "Flag state",
    "role": "Role on board",
    "employer": "Employer / crewing agency",
    "shipowner": "Shipowner / operator",
    "contractTerms": "Employment agreement / CBA",
}
CASE_DATA_KEYS = (
    "situation",
    "shipName",
    "imoNumber",
    "flagState",
    "role",
    "employer",
    "shipowner",
    "incidentDate",
    "incidentLocation",
    "medicalStatus",
    "contractTerms",
    "evidence",
    "communications",
    "desiredOutcome",
)


def int_setting(name: str, default: int, *, low: int, high: int) -> int:
    try:
        value = int(get_setting(name, default))
    except (TypeError, ValueError):
        value = default
    return max(low, min(value, high))


def build_config() -> ProviderConfig:
    try:
        timeout = int(get_setting("REQUEST_TIMEOUT_SECONDS", 420))
    except (TypeError, ValueError):
        timeout = 420
    effort = str(get_setting("ANTHROPIC_EFFORT", "high") or "high").strip().lower()
    if effort not in {"low", "medium", "high", "xhigh", "max"}:
        effort = "high"
    settings: dict[str, Any] = {
        "openai_key": str(get_setting("OPENAI_API_KEY", "") or ""),
        "anthropic_key": str(get_setting("ANTHROPIC_API_KEY", "") or ""),
        "zai_key": str(get_setting("ZAI_API_KEY", "") or ""),
        "openai_model": str(get_setting("OPENAI_MODEL", "gpt-5.6-sol") or "gpt-5.6-sol"),
        "anthropic_model": str(
            get_setting("ANTHROPIC_MODEL", "claude-opus-5") or "claude-opus-5"
        ),
        "zai_model": str(get_setting("ZAI_MODEL", "glm-5.3") or "glm-5.3"),
        "demo_mode": DEMO_MODE,
        "timeout_seconds": max(60, min(timeout, 900)),
        "anthropic_max_tokens": int_setting(
            "ANTHROPIC_MAX_TOKENS", 32_000, low=4_000, high=128_000
        ),
        "anthropic_effort": effort,
        "zai_max_tokens": int_setting(
            "ZAI_MAX_TOKENS", 32_000, low=4_000, high=128_000
        ),
    }

    # Streamlit re-executes this script on every rerun but keeps imported
    # modules in sys.modules, so a redeployed app can run a new page against a
    # mariner_core left over from the previous version. Passing a setting that
    # older core does not define raises a TypeError that Streamlit Cloud
    # redacts, which is close to undiagnosable. Detect the mismatch instead and
    # tell the operator the one thing that fixes it.
    supported = {field.name for field in fields(ProviderConfig)}
    st.session_state.stale_core_settings = sorted(set(settings) - supported)
    return ProviderConfig(
        **{key: value for key, value in settings.items() if key in supported}
    )


def render_stale_core_notice() -> None:
    stale = st.session_state.get("stale_core_settings") or []
    if not stale:
        return
    st.error(
        "This app is running a newer page against an older copy of `mariner_core.py` "
        "still loaded in memory, so these settings were ignored: "
        + ", ".join(stale)
        + ". Provider fixes that depend on them are **not active**. Reboot the app "
        "(Manage app → ⋮ → Reboot app) to load the current code."
    )


def require_shared_password() -> None:
    expected = str(get_setting("APP_PASSWORD", "") or "")
    if not expected or st.session_state.get("authenticated"):
        return
    st.html(STATIC_CSS)
    st.title("⚓ Mariner Advocate")
    st.caption("Private workspace access")
    entered = st.text_input("Shared password", type="password", autocomplete="current-password")
    if st.button("Open workspace", type="primary"):
        if hmac.compare_digest(entered.encode("utf-8"), expected.encode("utf-8")):
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("That password is not correct.")
    st.info(
        "This shared-password gate is a basic safeguard, not enterprise authentication. "
        "A real deployment still needs TLS, access controls, logging safeguards, and a privacy review."
    )
    st.stop()


def collect_case_data() -> dict[str, str]:
    return {
        key: str(st.session_state.get(f"case_{key}", "") or "")
        for key in CASE_DATA_KEYS
    }


def clear_session_callback() -> None:
    authenticated = st.session_state.get("authenticated")
    next_upload_epoch = int(st.session_state.get("contract_upload_epoch", 0) or 0) + 1
    for key in list(st.session_state):
        del st.session_state[key]
    if authenticated:
        st.session_state.authenticated = authenticated
    st.session_state.contract_upload_epoch = next_upload_epoch
    for key in CASE_DATA_KEYS:
        st.session_state[f"case_{key}"] = ""


def invalidate_case_results() -> None:
    for stale_key in ("panel", "submitted_case", "draft", "discussion"):
        st.session_state.pop(stale_key, None)


def extract_uploaded_contract_text(uploaded_file: Any) -> str:
    name = str(getattr(uploaded_file, "name", "contract") or "contract")
    suffix = Path(name).suffix.lower()
    raw = uploaded_file.getvalue()
    if not raw:
        raise AppError("The uploaded contract is empty.")
    if len(raw) > MAX_CONTRACT_UPLOAD_BYTES:
        raise AppError("The contract is larger than 10 MB. Upload a smaller copy.")

    try:
        if suffix == ".pdf":
            if not raw.lstrip().startswith(b"%PDF-"):
                raise AppError("The uploaded file is not a valid PDF.")
            if re.search(rb"/(?:JavaScript|JS|Launch|EmbeddedFile)\b", raw):
                raise AppError(
                    "PDFs containing scripts, launch actions, or embedded files are not accepted."
                )
            reader = PdfReader(BytesIO(raw))
            if reader.is_encrypted:
                raise AppError("The PDF is password-protected. Upload an unlocked copy.")
            if len(reader.pages) > MAX_CONTRACT_PDF_PAGES:
                raise AppError("The PDF has more than 150 pages. Upload only the relevant contract.")
            parts = []
            for page_number, page in enumerate(reader.pages, start=1):
                page_text = str(page.extract_text() or "").strip()
                if page_text:
                    parts.append(f"[Page {page_number}]\n{page_text}")
            text = "\n\n".join(parts)
        elif suffix == ".docx":
            if not zipfile.is_zipfile(BytesIO(raw)):
                raise AppError("The uploaded file is not a valid DOCX document.")
            with zipfile.ZipFile(BytesIO(raw)) as archive:
                entries = archive.infolist()
                names = {entry.filename for entry in entries}
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    raise AppError("The uploaded file is not a valid DOCX document.")
                if any(name.lower().endswith("vbaproject.bin") for name in names):
                    raise AppError("Macro-enabled Word documents are not accepted.")
                if len(entries) > MAX_DOCX_ENTRIES:
                    raise AppError("The DOCX contains too many embedded entries.")
                if any(
                    entry.filename.startswith(("/", "\\"))
                    or ".." in PurePosixPath(entry.filename).parts
                    for entry in entries
                ):
                    raise AppError("The DOCX contains unsafe internal paths.")
                uncompressed_size = sum(entry.file_size for entry in entries)
                if uncompressed_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise AppError("The expanded DOCX is too large to process safely.")
                if any(
                    entry.file_size > 1_000_000
                    and entry.file_size / max(entry.compress_size, 1) > 200
                    for entry in entries
                ):
                    raise AppError("The DOCX uses an unsafe compression ratio.")
            document = Document(BytesIO(raw))
            parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            for table_number, table in enumerate(document.tables, start=1):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    parts.append(f"[Table {table_number}]\n" + "\n".join(rows))
            text = "\n\n".join(parts)
        elif suffix == ".txt":
            if b"\x00" in raw:
                raise AppError("The TXT upload appears to contain binary data.")
            control_bytes = sum(byte < 32 and byte not in (9, 10, 13) for byte in raw)
            if control_bytes > max(5, len(raw) // 100):
                raise AppError("The TXT upload appears to contain binary data.")
            try:
                text = raw.decode("utf-8-sig")
            except UnicodeDecodeError:
                text = raw.decode("cp1252")
        else:
            raise AppError("Use a PDF, DOCX, or TXT contract file.")
    except AppError:
        raise
    except Exception as exc:
        raise AppError(
            "The contract could not be read. Try exporting it again as PDF, DOCX, or TXT."
        ) from exc

    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.replace("\x00", "").splitlines()]
    text = "\n".join(line for line in lines if line).strip()
    if len(text) < 40:
        raise AppError(
            "The contract has too little readable text. If it is a scan, run OCR and upload it again."
        )
    if len(text) > MAX_CONTRACT_CHARS:
        raise AppError(
            "The extracted contract is too long. Upload the employment agreement and relevant annexes only."
        )
    return text


def apply_contract_fields(
    extraction: dict[str, Any], *, overwrite: bool
) -> tuple[list[str], list[str]]:
    fields = extraction.get("fields") if isinstance(extraction.get("fields"), dict) else {}
    filled: list[str] = []
    preserved: list[str] = []
    for field, label in CONTRACT_CASE_FIELDS.items():
        value = str(fields.get(field) or "").strip()
        if not value:
            continue
        state_key = f"case_{field}"
        if str(st.session_state.get(state_key, "") or "").strip() and not overwrite:
            preserved.append(label)
            continue
        st.session_state[state_key] = value
        filled.append(label)
    if filled:
        invalidate_case_results()
    return filled, preserved


def render_contract_import_summary() -> None:
    extraction = st.session_state.get("contract_import_result")
    if not isinstance(extraction, dict):
        return
    filename = str(st.session_state.get("contract_import_filename") or "contract")
    with st.expander(f"Review last contract import: {filename}"):
        clauses = extraction.get("importantClauses") or []
        if clauses:
            st.markdown("**Relevant clauses found**")
            for clause in clauses:
                if not isinstance(clause, dict):
                    continue
                topic = str(clause.get("topic") or "Relevant clause")
                location = str(clause.get("pageOrSection") or "")
                summary = str(clause.get("summary") or clause.get("quote") or "")
                st.write(f"{topic}{f' — {location}' if location else ''}: {summary}")
        missing = [str(item) for item in extraction.get("missingFields") or [] if str(item)]
        if missing:
            st.caption("Not found or unclear: " + ", ".join(missing))
        for warning in extraction.get("warnings") or []:
            if str(warning):
                st.warning(str(warning))


def apply_pending_contract_import() -> None:
    pending = st.session_state.pop("pending_contract_import", None)
    if not isinstance(pending, dict):
        return
    extraction = pending.get("extraction")
    if not isinstance(extraction, dict):
        return
    filled, preserved = apply_contract_fields(
        extraction, overwrite=bool(pending.get("overwrite"))
    )
    st.session_state.contract_import_result = extraction
    st.session_state.contract_import_filename = Path(
        str(pending.get("filename") or "contract")
    ).name[:180]
    st.session_state.contract_import_notice = {
        "filled": filled,
        "preserved": preserved,
    }
    st.session_state.contract_upload_epoch = (
        int(st.session_state.get("contract_upload_epoch", 0) or 0) + 1
    )


def render_contract_importer(
    config: ProviderConfig,
) -> tuple[Any | None, bool, bool, bool]:
    st.markdown("#### Import an employment contract")
    with st.container(border=True):
        st.caption(
            "Upload a PDF, DOCX, or TXT copy of the SEA, employment agreement, or CBA. "
            "The importer fills only contract-related fields, and every value remains editable."
        )
        notice = st.session_state.pop("contract_import_notice", None)
        if isinstance(notice, dict):
            filled_notice = notice.get("filled") or []
            preserved_notice = notice.get("preserved") or []
            if filled_notice:
                st.success(
                    "Filled: "
                    + ", ".join(str(item) for item in filled_notice)
                    + ". Review before submitting."
                )
            else:
                st.warning("No empty case fields could be filled from this contract.")
            if preserved_notice:
                st.info(
                    "Preserved existing values: "
                    + ", ".join(str(item) for item in preserved_notice)
                    + "."
                )

        if PUBLIC_DEMO_ONLY:
            st.info(
                "Contract upload is disabled in the public fictional-data demo. "
                "Use a reviewed private deployment for document import."
            )
            render_contract_import_summary()
            return None, False, False, False

        uploaded_file = st.file_uploader(
            "Contract file",
            type=["pdf", "docx", "txt"],
            key=f"contract_upload_{int(st.session_state.get('contract_upload_epoch', 0) or 0)}",
            help="Maximum 10 MB. Scanned PDFs require searchable text/OCR.",
        )
        overwrite = st.checkbox(
            "Replace case fields that already contain text",
            value=False,
            key="contract_import_overwrite",
        )
        consent = True
        if not config.demo_mode:
            consent = st.checkbox(
                "I understand that the extracted contract text will be sent to OpenAI for field extraction.",
                key="contract_import_consent",
            )

        provider_ready = config.demo_mode or bool(config.openai_key)
        if not provider_ready:
            st.info("Add OPENAI_API_KEY to enable contract extraction.")

        import_clicked = st.form_submit_button(
            "Extract and fill contract fields",
            use_container_width=True,
            disabled=not provider_ready,
        )
        render_contract_import_summary()
        return uploaded_file, overwrite, consent, import_clicked


def safe_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def item_text(item: Any, key: str, default: str = "") -> str:
    if not isinstance(item, dict):
        return default
    return str(item.get(key) or default)


def dedupe(items: list[Any], key: str) -> list[dict[str, Any]]:
    seen: set[str] = set()
    result: list[dict[str, Any]] = []
    for item in items:
        if not isinstance(item, dict):
            continue
        marker = str(item.get(key) or "").strip().lower()
        if not marker or marker in seen:
            continue
        seen.add(marker)
        result.append(item)
    return result


def valid_web_url(value: str) -> bool:
    try:
        parsed = urlparse(value)
        return parsed.scheme in {"http", "https"} and bool(parsed.netloc)
    except ValueError:
        return False


def slug(value: str) -> str:
    cleaned = "".join(char.lower() if char.isalnum() else "-" for char in value)
    return "-".join(part for part in cleaned.split("-") if part)[:70] or "legal-draft"


def render_sidebar(config: ProviderConfig) -> None:
    with st.sidebar:
        st.header("Workspace")
        if PUBLIC_DEMO_ONLY:
            st.error("Public demo: fictional data only")
            st.caption("Provider calls are disabled and deterministic demo responses are used.")
        elif config.demo_mode:
            st.info("Demo mode is on. No provider calls will be made.")
        else:
            st.caption("Server-side provider status")
            statuses = {
                "OpenAI": bool(config.openai_key),
                "Anthropic": bool(config.anthropic_key),
                "Z.AI": bool(config.zai_key),
            }
            for name, ready in statuses.items():
                st.write(f"{'✅' if ready else '❌'} {name}")
        st.divider()
        st.caption("Configured models")

        def budget(name: str) -> str:
            # Tolerate a stale mariner_core still held in memory after a redeploy;
            # render_stale_core_notice explains the mismatch in the main column.
            value = getattr(config, name, None)
            return f"{value:,}" if isinstance(value, int) else "unknown"

        effort = getattr(config, "anthropic_effort", "unknown")
        st.code(
            f"Z.AI: {config.zai_model} (max {budget('zai_max_tokens')} out)\n"
            f"Claude: {config.anthropic_model} "
            f"({effort} effort, max {budget('anthropic_max_tokens')} out)\n"
            f"OpenAI: {config.openai_model}",
            language=None,
        )

        panel = st.session_state.get("panel")
        case_data = st.session_state.get("submitted_case")
        if isinstance(panel, dict) and isinstance(case_data, dict):
            export = json.dumps(
                {"caseData": case_data, "panel": panel}, ensure_ascii=False, indent=2
            )
            st.download_button(
                "Download case export",
                export,
                file_name="mariner-advocate-case.json",
                mime="application/json",
                use_container_width=True,
            )

        st.button(
            "Clear this session",
            use_container_width=True,
            on_click=clear_session_callback,
        )

        st.divider()
        st.caption(
            "API keys are read only from server environment variables or Streamlit secrets; "
            "they are never entered in this page or included in case exports."
        )


def render_header() -> None:
    st.markdown(
        """
        <div class="hero">
          <div class="eyebrow">LEGAL ISSUE-SPOTTING FOR SEAFARERS</div>
          <h1>A clearer path through a difficult moment.</h1>
          <p>Describe the incident once. Two models form independent views, a reviewer
          maps disagreements, and a senior model turns the supported result into practical
          next steps and careful document drafts.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if PUBLIC_DEMO_ONLY:
        st.error(
            "This public deployment is for fictional demonstrations only. Do not enter real "
            "names, medical facts, employment details, documents, or other personal data. "
            "Streamlit Community Cloud must not be used for this family's real case."
        )
    else:
        st.warning(
            "This supports, but does not replace, a qualified maritime lawyer, union representative, "
            "insurer, competent authority, or medical team. Do not delay treatment. Case facts are "
            "sent to the three configured AI providers when a live panel runs."
        )

    cards = [
        ("01", "Independent analysis", "Z.AI reviews facts and possible legal routes."),
        ("02", "Blind cross-check", "Claude commits its own view before seeing AI 1."),
        ("03", "Dispute review", "Claude compares the committed views and maps gaps."),
        ("04", "Senior arbitration", "OpenAI resolves supported disputes and flags uncertainty."),
    ]
    columns = st.columns(4)
    for column, (number, title, detail) in zip(columns, cards):
        with column:
            st.markdown(
                f'<div class="workflow-card"><span class="number">{number}</span>'
                f"<strong>{title}</strong><small>{detail}</small></div>",
                unsafe_allow_html=True,
            )


def render_intake(config: ProviderConfig) -> None:
    apply_pending_contract_import()
    missing: list[str] = []
    configured_provider_count = 3
    if not config.demo_mode:
        if not config.zai_key:
            missing.append("ZAI_API_KEY")
        if not config.anthropic_key:
            missing.append("ANTHROPIC_API_KEY")
        if not config.openai_key:
            missing.append("OPENAI_API_KEY")
        configured_provider_count = 3 - len(missing)
    review_disabled = not config.demo_mode and configured_provider_count < 2
    if missing:
        if review_disabled:
            st.error(
                "Live review needs at least two configured AI providers. Missing: "
                + ", ".join(missing)
            )
        else:
            st.warning(
                "The panel can run with reduced coverage. This provider is not configured: "
                + ", ".join(missing)
            )

    st.subheader("Case intake")
    st.caption("Use facts you know and mark uncertain details as uncertain.")
    with st.form("case_intake", border=True):
        (
            uploaded_contract,
            overwrite_contract_fields,
            contract_import_consent,
            import_submitted,
        ) = render_contract_importer(config)
        st.text_area(
            "What happened? *",
            height=220,
            key="case_situation",
            placeholder=(
                "Describe the incident in chronological order: what happened, where the person "
                "was, what work was underway, the injury, the medical response, and what the "
                "company has said."
            ),
        )

        st.markdown("#### Vessel and employment")
        col_a, col_b = st.columns(2)
        with col_a:
            st.text_input("Ship name", key="case_shipName")
            st.text_input("Flag state (high priority)", key="case_flagState")
            st.text_input("Employer / crewing agency", key="case_employer")
        with col_b:
            st.text_input("IMO number", key="case_imoNumber")
            st.text_input("Role on board", key="case_role")
            st.text_input("Shipowner / operator", key="case_shipowner")

        st.markdown("#### Incident and treatment")
        col_c, col_d = st.columns(2)
        with col_c:
            st.text_input("Incident date and time", key="case_incidentDate")
        with col_d:
            st.text_input("Location / port / waters", key="case_incidentLocation")
        st.text_area(
            "Injury and current treatment",
            height=110,
            key="case_medicalStatus",
            placeholder="Use clinicians' wording where possible; do not speculate.",
        )

        st.markdown("#### Documents and communications")
        st.text_area(
            "Employment agreement / CBA",
            height=90,
            key="case_contractTerms",
            placeholder="Paste relevant clauses or note which documents are available.",
        )
        st.text_area(
            "Incident report and evidence",
            height=90,
            key="case_evidence",
            placeholder="Logbook, witnesses, CCTV, photos, weather, risk assessment, medical report, receipts.",
        )
        st.text_area(
            "What has the company or insurer said?",
            height=90,
            key="case_communications",
            placeholder="Paste exact wording where possible, including dates and who said it.",
        )
        st.text_area(
            "Desired outcome",
            height=75,
            key="case_desiredOutcome",
            placeholder="Medical payment, sick wages, repatriation, disability benefit, evidence, complaint...",
        )

        with st.expander("Privacy checklist before submitting"):
            st.write(
                "Remove passport numbers, banking details, full home addresses, passwords, and "
                "unrelated medical history. Share only what is necessary. This app does not "
                "persist cases to a database, but live case facts are sent to the configured providers."
            )

        fictional_confirmed = True
        if PUBLIC_DEMO_ONLY:
            fictional_confirmed = st.checkbox(
                "I confirm that I am using entirely fictional data in this public demo.",
                key="fictional_data_confirmed",
            )

        panel_submitted = st.form_submit_button(
            "Run the legal review panel",
            type="primary",
            use_container_width=True,
            disabled=review_disabled,
        )

    if import_submitted:
        submitted_case = st.session_state.get("submitted_case")
        if isinstance(submitted_case, dict) and collect_case_data() != submitted_case:
            invalidate_case_results()
        if uploaded_contract is None:
            st.error("Choose a PDF, DOCX, or TXT contract before extracting fields.")
            return
        if not contract_import_consent:
            st.error("Confirm the OpenAI document-extraction notice before continuing.")
            return
        try:
            with st.spinner("Reading the contract and extracting stated terms..."):
                contract_text = extract_uploaded_contract_text(uploaded_contract)
                extraction = extract_contract_fields(contract_text, config)
            st.session_state.pending_contract_import = {
                "extraction": extraction,
                "overwrite": overwrite_contract_fields,
                "filename": str(uploaded_contract.name),
            }
            st.rerun()
        except AppError as exc:
            st.error(str(exc))
        except Exception:
            st.error("The contract import failed because of an unexpected server error.")
        return

    if not panel_submitted:
        return
    if not fictional_confirmed:
        st.error("Confirm that the public demo contains fictional data before continuing.")
        return

    case_data = collect_case_data()
    try:
        with st.status("Panel in session", expanded=True) as status:
            panel = analyze_case(case_data, config, progress=status.write)
            workflow_status = str((panel.get("workflow") or {}).get("status") or "failed")
            if workflow_status == "complete":
                status.update(label="Panel review complete", state="complete", expanded=False)
            elif workflow_status == "degraded":
                status.update(
                    label="Review finished with reduced provider coverage",
                    state="complete",
                    expanded=False,
                )
            elif workflow_status == "partial":
                status.update(
                    label="Senior arbitration failed; partial analyses were preserved",
                    state="error",
                    expanded=False,
                )
            else:
                status.update(
                    label="No usable provider panel could be completed",
                    state="error",
                    expanded=False,
                )
        st.session_state.panel = panel
        st.session_state.submitted_case = case_data
        st.session_state.pop("draft", None)
        st.session_state.pop("discussion", None)
        failures = safe_list(panel.get("providerFailures"))
        warnings = safe_list(panel.get("providerWarnings"))
        if workflow_status == "degraded" and failures:
            st.warning(
                f"Review finished with reduced coverage: {len(failures)} provider stage(s) "
                "were unavailable. The available models continued."
            )
        elif workflow_status == "degraded":
            st.warning(
                f"Review finished with reduced coverage: {len(warnings)} provider stage(s) "
                "were cut short by their output limit and only partly recovered."
            )
        elif workflow_status == "partial":
            st.error("Senior arbitration failed. Any completed earlier analyses were preserved.")
        elif workflow_status == "failed":
            st.error("No independent provider analysis completed. No arbitration was produced.")
        else:
            st.success("The panel completed its review.")
    except AppError as exc:
        st.error(str(exc))
    except Exception:
        st.error("The panel could not finish because of an unexpected server error.")


def render_overview(panel: dict[str, Any]) -> None:
    arbitration = panel.get("arbitration")
    workflow = panel.get("workflow") or {}
    if not isinstance(arbitration, dict):
        st.error(
            "Senior OpenAI arbitration was unavailable. Any completed earlier analyses "
            "are preserved in the Analyses tab, but there is no final synthesis."
        )
        return
    st.info(str(arbitration.get("executiveSummary") or "No synthesis was returned."))
    metric_a, metric_b, metric_c = st.columns(3)
    metric_a.metric("Overall confidence", str(arbitration.get("overallConfidence") or "unclear").title())
    metric_b.metric("Recommended actions", len(safe_list(arbitration.get("recommendedActions"))))
    metric_c.metric("Open questions", len(safe_list(arbitration.get("unresolvedQuestions"))))

    st.subheader("What to do next")
    actions = safe_list(arbitration.get("recommendedActions"))
    for index, action in enumerate(actions, start=1):
        if not isinstance(action, dict):
            continue
        order = action.get("order") or index
        title = item_text(action, "action", "Action")
        with st.container(border=True):
            st.markdown(f"**{order}. {title}**")
            details = [
                item_text(action, "purpose"),
                f"Owner: {item_text(action, 'owner')}" if item_text(action, "owner") else "",
                f"Timing: {item_text(action, 'timing')}" if item_text(action, "timing") else "",
            ]
            st.caption(" · ".join(detail for detail in details if detail))

    st.subheader("Possible rights and claims")
    for right in safe_list(arbitration.get("provisionalRights")):
        if not isinstance(right, dict):
            continue
        with st.expander(
            f"{item_text(right, 'rightOrClaim', 'Possible claim')} — {item_text(right, 'status', 'unclear')}"
        ):
            st.write(item_text(right, "basis", "Basis not returned."))
            st.caption("Verification needed: " + item_text(right, "verification", "Not specified."))

    warnings = safe_list(arbitration.get("doNotDoYet"))
    if warnings:
        st.subheader("Do not do these yet")
        for warning in warnings:
            if isinstance(warning, dict):
                st.error(f"{item_text(warning, 'action')} — {item_text(warning, 'reason')}")

    st.caption(
        str(workflow.get("anchoringControl") or "")
        + " This output is decision support, not legal representation."
    )


def render_panel_health(panel: dict[str, Any]) -> None:
    failures = safe_list(panel.get("providerFailures"))
    warnings = safe_list(panel.get("providerWarnings"))
    if not failures and not warnings:
        return

    workflow = panel.get("workflow") or {}
    status = str(workflow.get("status") or "degraded")
    if status == "degraded" and not failures:
        st.warning(
            "Reduced panel coverage: a provider ran out of output allowance mid-answer. "
            "The recovered part of its response was used; later sections are missing."
        )
    elif status == "degraded":
        st.warning(
            "Reduced panel coverage: at least one provider stage failed, but the available "
            "models continued and OpenAI completed a confidence-limited synthesis."
        )
    elif status == "partial":
        st.error(
            "Partial panel only: senior arbitration failed, but completed earlier analyses "
            "have been preserved below."
        )
    else:
        st.error("The provider panel could not produce a usable final review.")

    stage_labels = {
        "firstAnalysis": "Z.AI first analysis",
        "independentAnalysis": "Claude independent analysis",
        "critique": "Claude comparison",
        "arbitration": "OpenAI senior arbitration",
    }
    with st.expander("Provider status details"):
        for entry in list(failures) + list(warnings):
            if not isinstance(entry, dict):
                continue
            stage = str(entry.get("stage") or "provider stage")
            model = str(entry.get("model") or "unknown model")
            message = str(entry.get("message") or "No usable response was returned.")
            st.markdown(f"**{stage_labels.get(stage, stage)} · {model}**")
            st.caption(message)
            detail = str(entry.get("detail") or "")
            error_type = str(entry.get("errorType") or "")
            technical = " · ".join(part for part in [error_type, detail] if part)
            if technical:
                st.code(technical, language=None)


def render_available_analyses(panel: dict[str, Any]) -> None:
    entries = [
        ("Z.AI first analysis", "first"),
        ("Claude independent analysis", "independent"),
        ("Claude comparison and critique", "critique"),
    ]
    for title, key in entries:
        st.subheader(title)
        result = panel.get(key)
        if not isinstance(result, dict):
            st.warning("This stage was unavailable. Other completed provider results were preserved.")
            continue
        if result.get("responseTruncated"):
            st.warning(
                "This model hit its output limit mid-answer. The complete leading part of "
                "its response is shown below; later sections are missing."
            )
        summary = str(result.get("summary") or "").strip()
        if summary:
            st.info(summary)
        urgent_actions = safe_list(result.get("urgentActions"))
        if urgent_actions:
            st.markdown("**Urgent actions surfaced by this model**")
            for action in urgent_actions:
                if isinstance(action, dict):
                    st.write(f"- {item_text(action, 'action', 'Action to confirm')}")
        with st.expander("View full structured response"):
            st.json(result)


def render_disputes(panel: dict[str, Any]) -> None:
    critique = panel.get("critique") or {}
    arbitration = panel.get("arbitration") or {}
    critique_available = isinstance(panel.get("critique"), dict)
    if not critique_available:
        st.warning("Claude's comparison stage was unavailable; no cross-model disputes can be shown.")
    disagreements = safe_list(critique.get("disagreements"))
    if critique_available:
        st.subheader("Where the models disagreed")
        if not disagreements:
            st.info("No material disagreement was returned.")
    for disagreement in disagreements:
        if not isinstance(disagreement, dict):
            continue
        topic = item_text(disagreement, "topic", "Disagreement")
        materiality = item_text(disagreement, "materiality", "unclear")
        with st.expander(f"{topic} — {materiality} materiality", expanded=True):
            left, right = st.columns(2)
            with left:
                st.caption("First analyst")
                st.write(item_text(disagreement, "firstPosition", "Not returned."))
            with right:
                st.caption("Independent reviewer")
                st.write(item_text(disagreement, "reviewerPosition", "Not returned."))
            st.write("**Why:**", item_text(disagreement, "reason", "Not returned."))
            st.write(
                "**Evidence that may resolve it:**",
                item_text(disagreement, "evidenceThatWouldResolve", "Not returned."),
            )

    st.subheader("Senior resolutions")
    for resolution in safe_list(arbitration.get("resolvedDisputes")):
        if not isinstance(resolution, dict):
            continue
        with st.container(border=True):
            st.markdown(f"**{item_text(resolution, 'topic')}**")
            st.write(item_text(resolution, "resolution"))
            st.caption(
                f"{item_text(resolution, 'reason')} · Confidence: {item_text(resolution, 'confidence', 'unclear')}"
            )

    unsupported = safe_list(critique.get("unsupportedClaims"))
    if unsupported:
        st.subheader("Unsupported claims caught")
        for claim in unsupported:
            if isinstance(claim, dict):
                st.error(
                    f"{item_text(claim, 'claim')} — {item_text(claim, 'problem')} "
                    f"Correction: {item_text(claim, 'correction')}"
                )

    st.subheader("Questions no model should guess")
    for question in safe_list(arbitration.get("unresolvedQuestions")):
        if isinstance(question, dict):
            st.warning(
                f"{item_text(question, 'question')}\n\nNeeded: {item_text(question, 'neededEvidence')}\n\n"
                f"Why it matters: {item_text(question, 'consequence')}"
            )


def render_evidence(panel: dict[str, Any]) -> None:
    first = panel.get("first") or {}
    independent = panel.get("independent") or {}
    arbitration = panel.get("arbitration") or {}
    missing = dedupe(
        safe_list(first.get("missingFacts")) + safe_list(independent.get("missingFacts")),
        "question",
    )
    evidence = dedupe(
        safe_list(first.get("evidenceChecklist"))
        + safe_list(independent.get("evidenceChecklist")),
        "item",
    )

    st.subheader("Questions to answer")
    for item in missing:
        with st.container(border=True):
            st.markdown(
                f"**{item_text(item, 'question')}** · priority: {item_text(item, 'priority', 'unclear')}"
            )
            st.caption(item_text(item, "whyItMatters"))

    st.subheader("Evidence preservation checklist")
    panel_id = str(panel.get("generatedAt") or "panel")
    for index, item in enumerate(evidence):
        marker = hashlib.sha256(
            f"{panel_id}:{index}:{item_text(item, 'item')}".encode("utf-8")
        ).hexdigest()[:12]
        st.checkbox(item_text(item, "item", "Evidence item"), key=f"evidence_{marker}")
        details = " · ".join(
            value
            for value in [item_text(item, "reason"), item_text(item, "howToPreserve")]
            if value
        )
        if details:
            st.caption(details)

    sources: list[dict[str, Any]] = []
    for analysis in (first, independent):
        for issue in safe_list(analysis.get("issues")):
            if isinstance(issue, dict):
                sources.extend(safe_list(issue.get("sources")))
    sources = dedupe(sources, "url")
    st.subheader("Sources surfaced — verify with counsel")
    for source in sources:
        title = item_text(source, "title", "Source")
        url = item_text(source, "url")
        status = item_text(source, "status", "needs-verification")
        if valid_web_url(url):
            st.link_button(f"{title} ({status})", url)
        else:
            st.write(f"{title} — invalid or missing URL — {status}")

    triggers = safe_list(arbitration.get("lawyerEscalationTriggers"))
    if triggers:
        st.subheader("Get qualified help promptly if...")
        for trigger in triggers:
            st.error(str(trigger))


def render_documents(panel: dict[str, Any], case_data: dict[str, Any], config: ProviderConfig) -> None:
    arbitration = panel.get("arbitration")
    if not isinstance(arbitration, dict):
        st.warning("Document drafting is unavailable until senior arbitration succeeds.")
        return
    plan = safe_list(arbitration.get("documentPlan"))
    st.subheader("Recommended document set")
    for item in plan:
        if not isinstance(item, dict):
            continue
        with st.container(border=True):
            st.markdown(
                f"**{item_text(item, 'document', 'Document')} → {item_text(item, 'recipient', 'Recipient to confirm')}**"
            )
            st.write(item_text(item, "purpose"))
            st.caption("Inputs needed: " + ", ".join(map(str, safe_list(item.get("inputsNeeded")))))

    standard_types = [
        "Incident and benefits notification",
        "Medical payment guarantee request",
        "Evidence preservation request",
        "On-board complaint",
        "Claim chronology",
    ]
    document_types = []
    for item in plan:
        if isinstance(item, dict) and item_text(item, "document"):
            document_types.append(item_text(item, "document"))
    document_types = list(dict.fromkeys(document_types + standard_types))

    st.subheader("Create a careful first draft")
    with st.form("document_form", border=True):
        document_type = st.selectbox("Document type", document_types)
        recipient = st.text_input(
            "Recipient", placeholder="Company, shipowner, insurer, P&I correspondent, authority..."
        )
        language = st.selectbox("Language", ["English", "Romanian"])
        extra = st.text_area(
            "Extra drafting instructions",
            placeholder="Optional preferences only. Do not paste new sensitive facts here unless necessary.",
        )
        make_draft = st.form_submit_button(
            "Create reviewed draft", type="primary", use_container_width=True
        )

    if make_draft:
        try:
            with st.status("OpenAI is drafting from the arbitrated record...", expanded=True) as status:
                draft = draft_document(
                    case_data,
                    panel,
                    config,
                    document_type=document_type,
                    recipient=recipient,
                    language=language,
                    extra_instructions=extra,
                )
                status.update(label="Draft complete", state="complete", expanded=False)
            st.session_state.draft = draft
        except AppError as exc:
            st.error(str(exc))
        except Exception:
            st.error("The draft could not be completed because of an unexpected server error.")

    draft = st.session_state.get("draft")
    if not isinstance(draft, dict):
        return
    title = str(draft.get("title") or "Legal draft")
    st.markdown(f"### {title}")
    draft_text = st.text_area(
        "Draft text — editable",
        value=str(draft.get("documentText") or ""),
        height=520,
        key=f"draft_text_{hashlib.sha256(title.encode('utf-8')).hexdigest()[:10]}",
    )
    st.download_button(
        "Download draft as .txt",
        draft_text,
        file_name=f"{slug(title)}.txt",
        mime="text/plain",
    )
    attachments = safe_list(draft.get("attachmentsChecklist"))
    fields = safe_list(draft.get("fieldsToConfirm"))
    if attachments:
        st.write("**Attachments checklist:**", ", ".join(map(str, attachments)))
    if fields:
        st.write("**Confirm before use:**", ", ".join(map(str, fields)))
    st.warning(
        str(draft.get("reviewWarning") or "Have a qualified adviser review this draft before use.")
    )


def render_followup(panel: dict[str, Any], case_data: dict[str, Any], config: ProviderConfig) -> None:
    if not isinstance(panel.get("arbitration"), dict):
        st.warning("Next-step discussion is unavailable until senior arbitration succeeds.")
        return
    st.subheader("Discuss the next required steps")
    st.caption(
        "Answers use the arbitrated record. They cannot verify new law, deadlines, medical facts, or documents."
    )
    history = st.session_state.setdefault("discussion", [])
    for turn in history:
        if not isinstance(turn, dict):
            continue
        with st.chat_message("user"):
            st.write(str(turn.get("question") or ""))
        response = turn.get("response") or {}
        with st.chat_message("assistant"):
            st.write(str(response.get("answer") or "No answer returned."))
            actions = safe_list(response.get("nextActions"))
            if actions:
                st.write("**Suggested sequence**")
                for action in actions:
                    if isinstance(action, dict):
                        st.write(
                            f"- {item_text(action, 'action')} — {item_text(action, 'owner')} — {item_text(action, 'timing')}"
                        )
            unknowns = safe_list(response.get("unknowns"))
            if unknowns:
                st.caption("Still unknown: " + ", ".join(map(str, unknowns)))
            if response.get("reviewWarning"):
                st.warning(str(response["reviewWarning"]))

    with st.form("followup_form", border=True):
        question = st.text_area(
            "Your question",
            placeholder="For example: What should we request from the company today?",
            height=90,
        )
        asked = st.form_submit_button("Ask about next steps", type="primary")
    if asked:
        try:
            with st.spinner("Reviewing the arbitrated record..."):
                response = answer_followup(
                    case_data,
                    panel,
                    config,
                    question=question,
                    history=history,
                )
            history.append({"question": question.strip(), "response": response})
            st.rerun()
        except AppError as exc:
            st.error(str(exc))
        except Exception:
            st.error("The follow-up could not be completed because of an unexpected server error.")

    if history and st.button("Clear discussion"):
        st.session_state.discussion = []
        st.rerun()


def render_results(config: ProviderConfig) -> None:
    panel = st.session_state.get("panel")
    case_data = st.session_state.get("submitted_case")
    if not isinstance(panel, dict) or not isinstance(case_data, dict):
        st.markdown(
            """
            <div class="case-empty">
              <div style="font-size:2rem">✦</div>
              <h3>Your review will appear here</h3>
              <p>The panel will separate facts from assumptions, compare legal views,
              identify missing evidence, and propose careful documents.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )
        return

    workflow_status = str((panel.get("workflow") or {}).get("status") or "failed")
    if isinstance(panel.get("arbitration"), dict):
        st.subheader("Arbitrated case review")
    elif workflow_status == "partial":
        st.subheader("Partial panel review")
    else:
        st.subheader("Provider panel report")
    render_panel_health(panel)
    tabs = st.tabs(
        ["Overview", "Analyses", "Disputes", "Evidence", "Documents", "Ask next steps"]
    )
    with tabs[0]:
        render_overview(panel)
    with tabs[1]:
        render_available_analyses(panel)
    with tabs[2]:
        render_disputes(panel)
    with tabs[3]:
        render_evidence(panel)
    with tabs[4]:
        render_documents(panel, case_data, config)
    with tabs[5]:
        render_followup(panel, case_data, config)


def main() -> None:
    require_shared_password()
    config = build_config()
    st.html(STATIC_CSS)
    render_sidebar(config)
    render_header()
    render_stale_core_notice()
    if not PUBLIC_DEMO_ONLY and not config.demo_mode and not str(get_setting("APP_PASSWORD", "") or ""):
        st.error(
            "Live mode is locked because APP_PASSWORD is not configured. Add a long random "
            "password in server secrets, then reload the app."
        )
        st.stop()
    intake, results = st.columns([0.92, 1.08], gap="large")
    with intake:
        render_intake(config)
    with results:
        render_results(config)
    st.divider()
    st.caption(
        "Mariner Advocate is decision support, not legal representation. No database is included. "
        "Review provider retention and data-processing terms before live use."
    )


if __name__ == "__main__":
    main()
